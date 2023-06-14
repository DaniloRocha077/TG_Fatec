"""This is a docstring which describes the module"""
import re
import streamlit as st
import docx
import pydantic
import connection.database as condb

# Configurações da página
st.set_page_config(page_title="Enviar TCC", page_icon="📤")

st.write("""
# Submeter seu trabalho para a Fatec Mogi Mirim
Aqui você pode submeter o seu trabalho no formato de Word (.docx) para validação dos itens obrigatórios do seu trabalho
""")

# Funções de extração
def extract_tema(doc):
    """
    Extrai o nome dos integrantes do TCC e o tema do trabalho.
    """
    encontrou_mogi_mirim = False
    for i, p in enumerate(doc.paragraphs):
        if "MOGI MIRIM" in p.text:
            encontrou_mogi_mirim = True
        elif encontrou_mogi_mirim and p.text.strip():
            integrante = p.text
            j = i + 1
            while j < len(doc.paragraphs) and doc.paragraphs[j].text.strip():
                integrante += " " + doc.paragraphs[j].text
                j += 1
            tema = ""
            if j < len(doc.paragraphs):
                j += 1
                while j < len(doc.paragraphs) and not doc.paragraphs[j].text.strip():
                    j += 1
                if j < len(doc.paragraphs):
                    tema = doc.paragraphs[j].text.strip()
            return tema
    return None, None
   
def extract_keywords(doc):
    """
    Extrai as palavras-chave do documento.
    """
    for p in doc.paragraphs:
        if "Palavras-chave:" in p.text:
            match = re.search(r'Palavras-chave:(.*)', p.text)
            if match:
                keywords = match.group(1).split(',')
                return [kw.strip() for kw in keywords]
    return []

def extract_city(doc):
    """
    Extrai o nome da cidade imediatamente anterior ao parágrafo que contém o ano.
    """
    city = ""
    for p in doc.paragraphs:
        if re.search(r'\d{4}', p.text):
            match = re.search(r'(\w+( \w+)*)\s*$', city)
            if match:
                return match.group(1)
            else:
                return ""
        else:
            city = p.text
    return ""

def extract_year(doc):
    """
    Extrai o primeiro número de quatro dígitos consecutivos do documento.
    """
    for p in doc.paragraphs:
        match = re.search(r'\d{4}', p.text)
        if match:
            return match.group(0)
    return None


def extract_autores(doc):
    """
    Extrai o nome dos integrantes do TCC.
    """
    encontrou_mogi_mirim = False
    for i, p in enumerate(doc.paragraphs):
        if "MOGI MIRIM" in p.text:
            encontrou_mogi_mirim = True
        elif encontrou_mogi_mirim and p.text.strip():
            integrante = p.text
            j = i + 1
            while j < len(doc.paragraphs) and doc.paragraphs[j].text.strip():
                integrante += " " + doc.paragraphs[j].text
                j += 1
            return integrante.strip()
    return None

def extract_orientador(doc):
    """
    Extrai o nome do orientador(a) do TCC.
    """
    for p in doc.paragraphs:
        match = re.search(r'Orientador(?:a)?:\s*(.*)', p.text, re.IGNORECASE)
        if match:
            orientador = match.group(1).split(',')
            return [o.strip() for o in orientador if o.strip()]
    return []

def extract_resumo(doc, section_title, next_section_title):
    """
    Extrai o texto de uma seção do TCC.
    """
    section_started = False
    section_ended = False
    section_text = ""

    for p in doc.paragraphs:
        if section_title.lower() in p.text.lower():
            section_started = True
        elif next_section_title.lower() in p.text.lower():
            section_ended = True

        if section_started and not section_ended:
            section_text += p.text.strip()
            section_text += " "

    section_text = re.sub(section_title, '', section_text, flags=re.IGNORECASE)
    section_text = section_text.strip()
    return section_text

def extract_introducao(doc, section_title, next_section_title):
    """
    Extrai o texto de uma seção do TCC.
    """
    section_started = False
    section_ended = False
    section_text = ""

    for p in doc.paragraphs:
        if section_title.lower() in p.text.lower():
            section_started = True
        elif next_section_title.lower() in p.text.lower():
            section_ended = True

        if section_started and not section_ended:
            section_text += p.text.strip()
            section_text += " "

    section_text = re.sub(section_title, '', section_text, flags=re.IGNORECASE)
    section_text = section_text.strip()
    return section_text

def extract_conclusao(doc, section_title, next_section_title):
    """
    Extrai o texto de uma seção do TCC.
    """
    section_started = False
    section_ended = False
    section_text = ""

    for p in doc.paragraphs:
        if section_title.lower() in p.text.lower():
            section_started = True
        elif next_section_title.lower() in p.text.lower():
            section_ended = True

        if section_started and not section_ended:
            section_text += p.text.strip()
            section_text += " "

    section_text = re.sub(section_title, '', section_text, flags=re.IGNORECASE)
    section_text = section_text.strip()
    return section_text
# UPLOAD DO ARQUIVO E EXTRAÇÃO
uploaded_file = st.file_uploader("Envie o seu arquivo", type=['doc', 'docx'])
if uploaded_file is not None:
    # Ler o arquivo
    doc = docx.Document(uploaded_file)

    # Extrair informações do documento
    autores = extract_autores(doc)
    orientador = extract_orientador(doc)
    tema = extract_tema(doc)
    cidade = extract_city(doc)
    ano = extract_year(doc)
    resumo = extract_resumo(doc, "RESUMO", "PALAVRAS-CHAVE")
    keywords = extract_keywords(doc)
    introducao = extract_introducao(doc, "INTRODUÇÃO", "Estrutura")
    conclusao = extract_conclusao(doc, "CONSIDERAÇÕES FINAIS", "REFERÊNCIAS")
    intro = condb.remover_caracteres_invalidos(introducao)

    #Usa esta variável para definir se todos os itens estão preenchidos
    enviar = True

    st.write("Confira os itens que foram identificados ou não")
    if tema:
        st.checkbox("Tema", value=True, disabled=True)
    else:
        enviar = False
        st.checkbox("Tema", value=False, disabled=True)
        st.error("Seu Trabalho não apresenta este item ou não foi encontrado, verifique a formatação padrão (para títulos de seções é obrigatório estar maiúsculo!")
    if ano:
        st.checkbox("Ano", value=True, disabled=True)
    else:
        enviar = False
        st.checkbox("Ano", value=False, disabled=True)
        st.error("Seu Trabalho não apresenta este item ou não foi encontrado, verifique a formatação padrão (para títulos de seções é obrigatório estar maiúsculo!")

    if cidade:
        st.checkbox("Cidade", value=True, disabled=True)
    else:
        enviar = False
        st.checkbox("Cidade", value=False, disabled=True)
        st.error("Seu Trabalho não apresenta este item ou não foi encontrado, verifique a formatação padrão (para títulos de seções é obrigatório estar maiúsculo!")

    if keywords:
        st.checkbox("Palavras-chave", value=True, disabled=True)
    else:
        enviar = False
        st.checkbox("Palavras-chave", value=False, disabled=True)
        st.error("Seu Trabalho não apresenta este item ou não foi encontrado, verifique a formatação padrão (para títulos de seções é obrigatório estar maiúsculo!")

    if autores:
        st.checkbox("Nome dos autores", value=True, disabled=True)
    else:
        enviar = False
        st.checkbox("Nome dos autores", value=False, disabled=True)
        st.error("Seu Trabalho não apresenta este item ou não foi encontrado, verifique a formatação padrão (para títulos de seções é obrigatório estar maiúsculo!")

    if orientador:
        st.checkbox("Orientador(a)", value=True, disabled=True)
    else:
        enviar = False
        st.checkbox("Orientador(a)", value=False, disabled=True)
        st.error("Seu Trabalho não apresenta este item ou não foi encontrado, verifique a formatação padrão (para títulos de seções é obrigatório estar maiúsculo!")

    if resumo:
        st.checkbox("Resumo", value=True, disabled=True)
    else:
        enviar = False
        st.checkbox("Resumo", value=False, disabled=True)
        st.error("Seu Trabalho não apresenta este item ou não foi encontrado, verifique a formatação padrão (para títulos de seções é obrigatório estar maiúsculo!")

    if introducao:
        st.checkbox("Introdução", value=True, disabled=True)
    else:
        enviar = False
        st.checkbox("Introdução", value=False, disabled=True)
        st.error("Seu Trabalho não apresenta este item ou não foi encontrado, verifique a formatação padrão (para títulos de seções é obrigatório estar maiúsculo!")

    if conclusao:
        st.checkbox("Considerações Finais", value=True, disabled=True)
    else:
        enviar = False
        st.checkbox("Considerações Finais", value=False, disabled=True)
        st.error("Seu Trabalho não apresenta este item ou não foi encontrado, verifique a formatação padrão (para títulos de seções é obrigatório estar maiúsculo!")
    
    col1, col2 = st.columns(2)
    with col1:
        st.write("Verifique se as informações estão corretas, após, clique em 'SUBMETER'")
    with col2:
        botao = st.button(label='SUBMETER')

    tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs(["Nome dos Autores", "Orientador(a)", "Tema", "Cidade", "Ano", "Resumo", "Palavra-chave", "Introdução"])
    with tab1:
        if autores:
            st.write(autores)
        else:
            st.write("Nome dos autores não encontrado.")
    with tab2:
        if orientador:
            st.write(", ".join(orientador))
        else:
            st.write("Orientador(a) não encontrado.")
    with tab3:
        if tema:
            st.write(tema)
        else:
            st.write("Tema não encontrado")
    with tab4:
        if cidade:
            st.write(cidade)
        else:
            st.write("Cidade não encontrada")
    with tab5:
        if ano:
            st.write(ano)
        else:
            st.write("Ano não encontrado")
    with tab6:
        if resumo:
            class TextoJustificado(pydantic.BaseModel):
                text: str
            text = resumo
            justified_text = f'<div style="text-align: justify">{text}</div>'
            st.markdown(justified_text, unsafe_allow_html=True)
        else:
            st.write("Resumo não encontrado")
    with tab7:
        if keywords:
            st.write(", ".join(keywords))
        else:
            st.write("Palavras-chave não encontradas.")
    with tab8:
        if introducao:
            class TextoJustificado(pydantic.BaseModel):
                text: str
            texto = introducao
            justified_text = f'<div style="text-align: justify">{texto}</div>'
            st.markdown(justified_text, unsafe_allow_html=True)
        else:
            st.error("Introdução não encontrada")
            
    if botao:
        if enviar:
            #Cria a conexão com o MySql
            db = condb.create_server_connection("db4free.net", 3306, "fatecmm", "DanRocha!@#qaz")
            st.write("Conectando ao Servidor")
            #Verifica se o banco existe
            verifica = condb.verificar_banco()
            st.write("Verificando acesso!")
            #Verifica se existe ou Cria o banco trabalhos_tg
            dbb = condb.banco_connection("db4free.net", 3306, "fatecmm", "DanRocha!@#qaz", "envios_tg")
            
            #Verifica se já existe ou Cria a tabela
            verificar_tb = condb.verificar_tabela()
            st.write("Enviando os Dados!")
            # Verifica se o trabalho já existe na base de dados
            cursor = dbb.cursor()
            cursor.execute(f"SELECT * FROM tbl_trabalhos WHERE tema='{tema}' AND autores='{autores}' AND resumo='{resumo}'")
            result = cursor.fetchone()
            if result is not None:
                st.markdown("# Já existe um trabalho com esses dados no banco de dados! 😟")
            else:
                #Envia os dados dos trabalho para o banco
                enviar_banco = condb.enviar_banco(autores, orientador, tema, cidade, ano, resumo, keywords, introducao, conclusao)
                st.markdown("# Trabalho enviado com sucesso! 😃")
        else:
            st.markdown("# Seu trabalho não contém todos os itens para envio! 😟")
else:
    st.write("Você não enviou o arquivo ainda")
